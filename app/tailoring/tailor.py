"""Resume tailoring: given JD + master resume, generate a tailored version.

Strategy:
1. Claude reads JD and identifies the top 5-7 keywords/skills the ATS will scan for.
2. Claude rewrites bullets to surface relevant experience first, swap synonyms
   to match JD vocabulary, and reorder sections by relevance.
3. Generate a 200-word cover letter referencing specific JD signals.
4. Render to .docx using the master template (preserves formatting/ATS friendliness).

This is intentionally a thin skeleton — the quality lives in the prompt, which
you should iterate on against your actual resume. Run a few tailoring jobs
manually, eyeball the outputs, refine.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Tuple, Optional

from docx import Document
from sqlmodel import select

from app.config import settings
from app.db.init_db import get_session
from app.db.models import Application, ApplicationStatus, Job
from app.qa_store.resolver import QAResolver

log = logging.getLogger(__name__)


def _register_llm_spend(kind: str) -> None:
    """Count a tailoring call against the process-wide LLM budget counter.

    Tailoring is Sonnet spend — several times pricier per call than a scoring
    final — and it used to be completely invisible to the budget: `/api/debug`
    and the caps reported only what the scoring lane had spent, so the number
    an operator read was not the number being billed. Registering here makes
    the counter truthful.

    Deliberately does NOT block. `llm_budget_exhausted()` gates *background*
    scoring, which is speculative work nobody asked for; tailoring is a paying
    user pressing a button, and failing it because a background lane drained
    the day's budget would be the wrong trade. The per-user ceiling
    (`tailor_abuse_daily_cap`) is what bounds tailoring — this only makes the
    spend observable.
    """
    try:
        from app.matching.reranker import _register_final_call
        _register_final_call()
    except Exception:  # accounting must never fail a user's tailor
        log.debug("tailor spend registration failed (%s)", kind, exc_info=True)

# Initialize canonical QA Resolver
qa_resolver = QAResolver()

TAILOR_SYSTEM = """You are a senior engineering-career editor. You make a candidate's REAL resume fit a specific role by EDITING it — not by rewriting it into a generic "AI resume." The output must read like the candidate wrote it after a focused editing pass, and must survive BOTH an ATS and a skeptical human reviewer who is actively screening for AI-generated applications (many now reject anything that reads machine-written).

FIRST PRINCIPLE — RESTRAINT (this overrides everything below):
- Preserve the candidate's own wording, sentence rhythm, and voice. Keep bullets that already work essentially as-is. Change only what genuinely improves fit for THIS specific job.
- Aim to touch roughly a third of the content, not all of it. A resume that is fully rewritten reads as machine-generated and gets rejected.
- When in doubt, change less. Reordering and light edits beat wholesale rewriting.

1. STRUCTURE & VOICE:
   - Output VALID markdown only. Preserve the exact section structure of the input resume.
   - Match the candidate's existing cadence: if their bullets are terse, stay terse; if they write in fragments, keep fragments. Do not "upgrade" their voice into polished corporate prose.

2. VARY STRUCTURE (anti-fingerprint — critical):
   - Real resumes are bursty and uneven. Deliberately VARY bullet length (some one line, some two) and openings. NOT every bullet should start with an action verb, and NOT every bullet needs a metric.
   - A uniform "[Verb] + tool + %number" on every single line is the #1 signature of AI writing. Break the pattern on purpose. Some bullets describe scope or ownership without a number — that is fine and more human.

3. BOLD SPARINGLY:
   - Bold at most 2-3 genuinely load-bearing technologies per section, where a human naturally would. Bolding every tool is an AI tell and hurts readability.

4. HONEST KEYWORD BRIDGING:
   - Never claim a skill the master resume doesn't support. If the JD needs something the candidate lacks, either weave it as genuinely adjacent experience in the candidate's OWN voice, or leave it out. Do not fabricate production experience.
   - Only add a short "currently learning / adjacent" note when it is natural and true for that person. Do NOT append a formulaic "Transitioning to / Adjacent Tools Under Study" block to every resume — that boilerplate is itself a fingerprint.
   - Do NOT invent jobs, degrees, dates, or metrics. Every number must appear verbatim in the master resume.

5. NO AI SLOP:
   - Ban filler and buzzwords: "leveraged", "synergized", "cutting-edge", "harnessing", "orchestrated seamless integrations", "state-of-the-art", "spearheaded", "drove efficiency", "revolutionized", "demonstrated expertise in", "passionate about", "results-driven", "proven track record".
   - Write like a working engineer: concrete, specific, occasionally plain. Prefer the candidate's real phrasing over anything that sounds optimized.

6. SKILLS SECTION:
   - Reorder to put JD-relevant skills first. Do not pad with skills the candidate doesn't have.

GOAL: the same person, edited for this role — never a new person generated for it."""

COVER_SYSTEM = """You write tight cover letters (180-220 words) using a Problem→Solution→Proof structure. They must NOT sound like generic cover letters.

The core structure (this is mandatory):
1. PROBLEM (opening, 1-2 sentences): Name the specific business/technical problem THIS company is hiring to solve, inferred from the JD. Frame it as "You're looking to <do X / solve Y>." Be concrete — pull the actual challenge from the responsibilities, not a generic statement.
2. SOLUTION (middle, 2-4 sentences): Map that problem to ONE specific thing the candidate has actually built or done. Use the formula: "I built/did <specific thing> at <company/project>." Name the real technologies and the real context from the resume.
3. PROOF (within the solution): Attach a concrete, quantified result from the resume — "which <reduced latency 24% / scaled to N users / cut release time 65%>." Only use metrics that appear in the resume. Never invent numbers.
4. WHY-THIS-COMPANY (1 sentence): One detail you'd only know if you actually cared about this company/role.
5. CLOSER (1 sentence): Forward-looking and specific. No "I look forward to hearing from you."

Hard rules:
- No "I am writing to apply for…" or "I am excited to…" openers. Open on THEIR problem, not your enthusiasm.
- Every claim must be grounded in the resume. Do NOT invent jobs, metrics, or technologies.
- Plain prose, no markdown, no bullet points. 180-220 words.
- Match the JD's vocabulary for key technologies (exact terms an ATS would scan).
- ANTI-FINGERPRINT: vary your sentence lengths and openings. Do not produce a rigid, identical skeleton that would look the same across many applications. Write in the candidate's plain voice — a real person who happens to be a good writer, not a template. Reviewers now screen for AI-written letters; uniform, over-polished structure is the tell."""


class Tailor:
    def __init__(self):
        self._anthropic_client = None
        self._openai_client = None
        self._active_backend = None
        
        # Shared process-wide clients (app/common/llm.py) — a fresh SDK client
        # per Tailor leaked an httpx pool + SSL context per tailor request.
        # Long-form generation needs more than the lane-safe default timeout;
        # with_options reuses the same connection pool.
        if settings.anthropic_api_key:
            try:
                from app.common.llm import shared_anthropic
                self._anthropic_client = shared_anthropic(timeout=180.0, max_retries=1)
                self._active_backend = "anthropic"
            except Exception:
                pass
        if settings.openai_api_key:
            try:
                from app.common.llm import shared_openai
                self._openai_client = shared_openai(timeout=180.0, max_retries=1)
                if not self._active_backend:
                    self._active_backend = "openai"
            except Exception:
                pass

    def _clean_tailored_resume(self, text: str) -> str:
        if not text:
            return ""
        # Strip markdown wrapper blocks
        text = re.sub(r"^```markdown\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*```$", "", text, flags=re.IGNORECASE)
        
        # Strip any leading horizontal rules (like prompt template's "---")
        lines = text.splitlines()
        while lines:
            first = lines[0].strip()
            if not first:
                lines.pop(0)
            elif re.match(r"^[-*_ ]+$", first) and len(first) >= 3:
                lines.pop(0)
            else:
                break
                
        # Strip any "CUSTOM HIGHLIGHTS" or "CRITICAL HIGHLIGHTS" sections appended at the end
        cleaned_lines = []
        skip = False
        for line in lines:
            stripped = line.strip().lower()
            if "custom highlights" in stripped or "critical framing" in stripped or "senior reviewer" in stripped:
                # Drop preceding horizontal rule if present
                if cleaned_lines and re.match(r"^[-*_ ]+$", cleaned_lines[-1].strip()) and len(cleaned_lines[-1].strip()) >= 3:
                    cleaned_lines.pop()
                skip = True
                continue
                
            if skip:
                continue
                
            cleaned_lines.append(line)
            
        return "\n".join(cleaned_lines).strip()

    def tailor_resume(self, master_resume_md: str, job: Job, variant: str = "variant_a",
                      custom_highlight_block: Optional[str] = None,
                      revision_notes: Optional[str] = None,
                      user_instruction: Optional[str] = None) -> str:
        # ── ATS exact-phrase targeting ──────────────────────────────────────
        # Find the JD phrases an ATS will scan for that are NOT already verbatim
        # in the master resume, so the tailor can incorporate them honestly.
        ats_block = ""
        try:
            from app.tailoring.ats_keywords import analyze as ats_analyze
            from app.tailoring.ats_keywords import skill_phrases
            ats = ats_analyze(job.description or "", master_resume_md)
            # Only phrases naming something a person can HAVE. The ranked list is
            # a frequency count over the JD, so it also carries title and
            # marketing fragments ("design and own", "ai platform company",
            # "backbone for enterprise"). Telling a generator to work those in
            # "using the EXACT phrasing" is how a JD responsibility line ends up
            # in someone's employment history — the very fabrication grounding
            # then has to catch and charge a rebuild for.
            targets = skill_phrases(ats.missing)
            if targets:
                ats_block = (
                    "\n\nATS PRIORITY PHRASES — these exact terms appear in the JD but are "
                    "MISSING verbatim from the resume. Where the candidate's real experience "
                    "supports it, incorporate the EXACT phrasing below (do not paraphrase, "
                    "do not invent experience). If the candidate has no real experience with "
                    "one of them, LEAVE IT OUT — a missing keyword costs a search hit, an "
                    "invented one costs the interview:\n"
                    + "\n".join(f"  - {p}" for p in targets)
                    + f"\n(Already covered: {', '.join(skill_phrases(ats.matched)[:8])})"
                )
        except Exception as e:
            log.warning("ATS keyword analysis failed (continuing without it): %s", e)

        highlights_block = ""
        if custom_highlight_block:
            highlights_block = (
                "\n\nCRITICAL FRAMING INSTRUCTIONS — prioritize and integrate these points "
                "honestly into the resume's summary, skills, or experience bullet points:\n"
                + custom_highlight_block.strip()
            )

        revision_block = ""
        if revision_notes:
            revision_block = (
                "\n\nQUALITY REVIEWER FEEDBACK — a previous draft of this tailored resume "
                "failed review for the reasons below. Fix every issue WITHOUT inventing "
                "experience; stay strictly grounded in the master resume:\n"
                + revision_notes.strip()[:2000]
            )

        user_block = ""
        if user_instruction and user_instruction.strip():
            user_block = (
                "\n\nCANDIDATE'S OWN DIRECTION (highest priority, but never at the cost of "
                "honesty — do not invent experience to satisfy it): "
                + user_instruction.strip()[:500]
            )

        prompt = f"""Job description:
---
Title: {job.title}
Company: {job.company}
{job.description[:5000]}
---{ats_block}{highlights_block}{revision_block}{user_block}

Return the tailored resume in markdown. No commentary.
Do NOT output the "CRITICAL FRAMING INSTRUCTIONS" or "CUSTOM HIGHLIGHTS" as a separate section in the tailored resume."""

        # A/B Testing routing
        use_openai = False
        if variant == "variant_b" and self._openai_client:
            use_openai = True
        elif not self._anthropic_client and self._openai_client:
            use_openai = True

        if not use_openai and self._active_backend == "anthropic" and self._anthropic_client:
            try:
                system = [
                    {"type": "text", "text": TAILOR_SYSTEM, "cache_control": {"type": "ephemeral"}},
                    {"type": "text", "text": f"Master resume (markdown):\n---\n{master_resume_md}\n---", "cache_control": {"type": "ephemeral"}},
                ]
                from app.common.llm import sampling
                resp = self._anthropic_client.messages.create(
                    model=settings.tailoring_model,
                    max_tokens=4000,
                    system=system,
                    messages=[{"role": "user", "content": prompt}],
                    **sampling(settings.tailoring_model, settings.tailoring_temperature),
                )
                _register_llm_spend("tailor_resume")
                return self._clean_tailored_resume(resp.content[0].text)
            except Exception as e:
                log.warning("Tailor: Anthropic failed, falling back to OpenAI: %s", e)

        if self._openai_client:
            resp = self._openai_client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=4000,
                temperature=settings.tailoring_temperature,
                messages=[
                    {"role": "system", "content": f"{TAILOR_SYSTEM}\n\nMaster resume (markdown):\n---\n{master_resume_md}\n---"},
                    {"role": "user", "content": prompt},
                ]
            )
            return self._clean_tailored_resume(resp.choices[0].message.content)

        raise RuntimeError("No LLM backend available for tailoring resume")

    def write_cover_letter(self, master_resume_md: str, job: Job, custom_highlight_block: Optional[str] = None) -> str:
        highlights_block = ""
        if custom_highlight_block:
            highlights_block = f"\n\nFocus areas to prioritize/highlight:\n{custom_highlight_block.strip()}"
        prompt = f"""Job:
Title: {job.title}
Company: {job.company}
{job.description[:4000]}
{highlights_block}

Write the cover letter using the Problem→Solution→Proof structure:
1. Open by naming the specific problem {job.company} is hiring to solve (infer it from the responsibilities above).
2. Map it to ONE thing the candidate actually built or did, with the real company/project and technologies.
3. Attach a concrete quantified result that appears in the resume.
4. One sentence on why {job.company} specifically.
5. One forward-looking closer.

Ground every claim in the resume. Invent nothing."""

        if self._active_backend == "anthropic" and self._anthropic_client:
            try:
                system = [
                    {"type": "text", "text": COVER_SYSTEM, "cache_control": {"type": "ephemeral"}},
                    {"type": "text", "text": f"Resume (markdown):\n{master_resume_md}", "cache_control": {"type": "ephemeral"}},
                ]
                from app.common.llm import sampling
                resp = self._anthropic_client.messages.create(
                    model=settings.cover_letter_model,
                    max_tokens=600,
                    system=system,
                    messages=[{"role": "user", "content": prompt}],
                    **sampling(settings.cover_letter_model, settings.tailoring_temperature),
                )
                _register_llm_spend("cover_letter")
                return resp.content[0].text
            except Exception as e:
                log.warning("Tailor: Anthropic failed for cover letter, falling back to OpenAI: %s", e)

        if self._openai_client:
            resp = self._openai_client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=600,
                temperature=settings.tailoring_temperature,
                messages=[
                    {"role": "system", "content": f"{COVER_SYSTEM}\n\nResume (markdown):\n{master_resume_md}"},
                    {"role": "user", "content": prompt},
                ]
            )
            return resp.choices[0].message.content

        raise RuntimeError("No LLM backend available for cover letter")


def _add_formatted_run(para, text: str) -> None:
    """Add text to a paragraph, rendering **bold** and *italic* markers."""
    parts = re.split(r"(\*\*[^*\n]+\*\*|\*[^*\n]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            para.add_run(part[1:-1]).italic = True
        else:
            para.add_run(part)


def _set_ats_safe_styles(doc) -> None:
    """Configure document styles for maximum ATS parse accuracy.

    Enterprise ATS (Workday, iCIMS, Taleo) parse DOCX by reading standard
    Word styles. Custom fonts, complex tables, text boxes, and headers/footers
    are often dropped entirely. This sets safe defaults:
    - Single-column layout (no tables, no text boxes)
    - Standard named styles (Normal, Heading 1/2/3, List Bullet)
    - 11pt Calibri — universally supported, ATS-safe font
    - Tight margins to maximize content density without overflow
    """
    from docx.shared import Pt, Inches, RGBColor

    # Margins: 0.75in all sides — tight but readable, avoids truncation
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    # Normal style: 11pt Calibri, single spacing, no extra space after paragraph
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.space_before = Pt(0)

    # Heading 1: Name / top-level — 14pt bold, dark, no space before
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(4)

    # Heading 2: Section headers (Experience, Education, Skills) — 12pt bold
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(2)

    # Heading 3: Job title / company lines — 11pt bold
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(1)

    # List Bullet: standard bullet style ATS parsers recognize
    try:
        lb = doc.styles["List Bullet"]
        lb.font.name = "Calibri"
        lb.font.size = Pt(11)
        lb.paragraph_format.space_after = Pt(1)
        lb.paragraph_format.left_indent = Inches(0.25)
    except Exception:
        pass


def _md_to_docx(md_text: str, out_path: Path) -> None:
    """Convert markdown resume to ATS-safe DOCX using standard Word styles.

    Produces a single-column document with no tables, text boxes, headers,
    or footers — the format most reliably parsed by enterprise ATS systems.
    """
    doc = Document()
    _set_ats_safe_styles(doc)

    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            # Suppress excessive blank lines — ATS parsers handle spacing via style
            continue
        if stripped.startswith("```"):
            continue
        if re.match(r"^[-*_ ]+$", stripped) and len(stripped) >= 3:
            continue
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_formatted_run(doc.add_paragraph(style="List Bullet"), stripped[2:])
        else:
            _add_formatted_run(doc.add_paragraph(style="Normal"), stripped)

    doc.save(out_path)


def _persist_tailored_to_storage(uid: str | None, application_id: int, files: list) -> None:
    """Best-effort copy of the tailored docs to Supabase Storage. The local
    ``data/tailored/`` dir is on the container's ephemeral disk, so a redeploy
    wipes it and "View Docs" breaks for older applications. Mirroring to the
    existing ``resume`` bucket (under a tailored/ prefix) lets the read path
    re-hydrate them. Fully non-fatal — a failure never blocks tailoring."""
    from app.config import settings
    if not (settings.use_supabase and uid and uid != "local"):
        return
    try:
        from app.db.supabase_client import service_client
        sb = service_client()
        for fp in files:
            p = Path(fp)
            if not p.exists():
                continue
            key = f"{uid}/tailored/app_{application_id}/{p.name}"
            try:
                sb.storage.from_("resume").upload(
                    key, p.read_bytes(), {"upsert": "true"})
            except Exception as _ue:
                log.debug("tailored upload failed for %s: %s", key, _ue)
    except Exception as e:
        log.debug("tailored storage persist skipped: %s", e)


def tailor_for_application(application_id: int, user_instruction: Optional[str] = None) -> Tuple[Path, Path]:
    """Generate tailored resume + cover letter for one application."""
    import random
    from datetime import datetime

    # --- Phase 1: read data in a short session, then close it ---
    with get_session() as session:
        app = session.get(Application, application_id)
        if not app:
            raise ValueError(f"Application {application_id} not found")
        job = session.get(Job, app.job_id)
        # snapshot all fields needed for LLM calls (detach-safe primitives)
        job_title = job.title
        job_company = job.company
        job_description = job.description
        job_url = job.url
        job_posted_at = job.posted_at
        profile_variant = app.profile_variant  # "backend" | "ai_agents" | "fullstack" | None
        custom_highlight_block = app.custom_highlight_block  # optional extra bullets from SeniorReviewer
        app_user_id = app.user_id  # owner — used to name output files after the real user

    # --- Phase 2: all LLM work outside any session (no lock held) ---
    variant = random.choice(["variant_a", "variant_b"])

    # Resolve the master resume. Prefer a SeniorReviewer-recommended profile
    # variant; otherwise use THIS user's own uploaded/synthesized resume
    # (multi-tenant) so tailoring is grounded in their real CV, not a shared one.
    # FAIL CLOSED on tenant identity. data/profiles/*.md are the FOUNDER's résumé
    # variants (real name, phone, email, LinkedIn, GitHub), committed to the repo
    # from the single-user era. Honouring a recommended variant for an arbitrary
    # tenant tailors THEIR application from the founder's CV — the founder's PII
    # is then what gets rendered into the résumé and cover letter they download
    # and send, and the grounding checker waves it through because the output is
    # faithfully grounded in that master. Only the founder/local dev may use a
    # variant; every other user is tailored from their own uploaded résumé.
    from app.common.tenancy import is_founder
    _variant_path = None
    if profile_variant:
        if is_founder(app_user_id):
            _variant_path = settings.profiles_dir / f"{profile_variant}.md"
        else:
            log.warning(
                "Tailor app %d: ignoring profile_variant=%s for non-founder user %s — "
                "tailoring from their own résumé (a variant file is another person's CV).",
                application_id, profile_variant, app_user_id)
    if _variant_path and _variant_path.exists():
        master = _variant_path.read_text(encoding="utf-8")
        profile_source = _variant_path.name
    else:
        from app.matching.pipeline import _load_resume
        master = _load_resume(user_id=app_user_id)
        profile_source = "user resume" if app_user_id else settings.resume_path.name

    log.info(
        "Tailoring app %d using profile=%s source=%s",
        application_id, profile_variant or "master", profile_source,
    )
    tailor = Tailor()

    # Build a lightweight Job-like object so Tailor methods still work
    class _JobSnapshot:
        title = job_title
        company = job_company
        description = job_description

    job_snap = _JobSnapshot()
    cover = tailor.write_cover_letter(master, job_snap, custom_highlight_block=custom_highlight_block)

    # ── L0: is there anything to gain? ──────────────────────────────────────
    # When the master already matches most of what this JD asks for, a full
    # regeneration cannot honestly add much — it can only re-word bullets that
    # already work, which costs a tailor, costs a verification pass, and gives
    # the model a fresh opportunity to fabricate. A measured audit found one
    # pairing at 72% master coverage where tailoring produced +0 phrases for
    # $0.066 and a grounding failure that forced a rebuild.
    #
    # This is a coverage bar, not a quality judgement: the résumé is already the
    # right résumé for this job, so we ship it as-is and say so.
    skipped_reason = None
    master_coverage = None
    try:
        from app.tailoring.ats_keywords import analyze as _ats_analyze
        master_coverage = _ats_analyze(job_description or "", master).coverage_pct
        _skip_bar = float(getattr(settings, "tailor_skip_coverage_pct", 0.0) or 0.0)
        if _skip_bar and master_coverage >= _skip_bar:
            skipped_reason = (
                f"Your résumé already covers {master_coverage:.0%} of this posting's "
                f"key terms, so it was sent as-is rather than rewritten. Rewriting a "
                f"résumé that already fits only makes it read machine-written."
            )
            log.info("Tailor app %d: L0 skip — master coverage %.0f%% >= %.0f%%",
                     application_id, master_coverage * 100, _skip_bar * 100)
    except Exception as _ce:
        log.debug("coverage pre-check skipped for app %d: %s", application_id, _ce)

    # ── Generate → quality-check → rebuild loop ─────────────────────────────
    # If grounding or the Resume Doctor fails a draft, regenerate ONCE with the
    # reviewer's issues fed back into the prompt instead of blocking the
    # application on the first bad draft. Best draft wins.
    MAX_TAILOR_ATTEMPTS = 1 if skipped_reason else 2
    resume_md = ""
    grounding_failed = False
    grounding_notes = None
    doctor_failed = False
    doctor_notes = None
    doctor_score = None
    doctor_ats = None          # ATS keyword coverage %
    doctor_verdict = None      # Haiku recruiter hiring signal (was log-only)
    doctor_weak = []           # weak bullets (missing verb/metric)
    doctor_banned = []         # banned/cliché words found
    doctor_integrity = []      # integrity issues (unbacked claims)
    doctor_human = None        # 0-100 "reads human" score (anti-fingerprint)
    doctor_fingerprints = []   # AI-writing tells detected
    human_failed = False       # anti-fingerprint gate — retries, never blocks
    locked_fields = []         # immutable sections restored verbatim from master
    fabrications = []          # deterministic fact additions (employer/date/…)
    bold_trimmed = 0           # inline emphasis removed after generation
    grounding_tier = None      # L0 / L1 / L2 / L3
    grounding_calls = 0        # batched verifier requests actually issued
    grounding_cache_hits = 0
    grounding_spans_changed = 0
    grounding_spans_verified = 0
    attempts_used = 0
    revision_notes = None

    for attempt in range(1, MAX_TAILOR_ATTEMPTS + 1):
        attempts_used = attempt
        if skipped_reason:
            # L0 — ship the master. It still goes through the lock layer,
            # grounding and the Doctor below, which is cheap: with nothing
            # rewritten, grounding reports L0 and makes no LLM call at all.
            resume_md = master
            break_after_checks = True
        else:
            break_after_checks = False
            resume_md = tailor.tailor_resume(
                master, job_snap, variant=variant,
                custom_highlight_block=custom_highlight_block,
                revision_notes=revision_notes,
                user_instruction=user_instruction,
            )

        # ── Lock layer ────────────────────────────────────────────────────────
        # Restore immutable factual sections (Education: degree/school/dates)
        # verbatim from the master BEFORE grounding/Doctor/render, so an altered
        # credential is structurally impossible rather than merely flagged.
        try:
            from app.tailoring.lock import restore_locked_fields
            resume_md, locked_fields = restore_locked_fields(resume_md, master)
            if locked_fields:
                log.info("Lock: restored immutable fields for app %d: %s",
                         application_id, ", ".join(locked_fields))
        except Exception as _le:
            log.warning("Lock layer skipped for app %d (non-fatal): %s", application_id, _le)

        # ── Style repair ──────────────────────────────────────────────────────
        # Cap inline emphasis deterministically. The prompt has asked for
        # sparing bold since it was written and the model produced 51 spans
        # anyway; removing emphasis changes no words and no facts, so this is
        # the one AI tell we can simply delete instead of asking about.
        try:
            from app.tailoring.style import trim_bold
            resume_md, bold_trimmed = trim_bold(resume_md)
            if bold_trimmed:
                log.info("Style: removed %d excess bold spans for app %d",
                         bold_trimmed, application_id)
        except Exception as _se:
            log.warning("Style repair skipped for app %d (non-fatal): %s", application_id, _se)

        # ── Deterministic fabrication guard ───────────────────────────────────
        # Set difference in the ADDITION direction over employers, held titles,
        # employment dates, degrees, institutions, certifications and numbers.
        # No model, no threshold, no similarity score: a fact in the output that
        # is not in the input was invented, and nothing about how plausible it
        # reads changes that. This runs BEFORE grounding because it is free and
        # because it catches the class of error grounding is weakest on — a
        # fabricated credential sitting in a section grounding does not treat as
        # a claim-bearing bullet.
        try:
            from app.tailoring.evidence import fabrication_violations
            fabrications = fabrication_violations(master, resume_md)
            if fabrications:
                log.warning("Fabrication guard: app %d added %d fact(s) not in the "
                            "master résumé: %s", application_id, len(fabrications),
                            "; ".join(f"{k}: {v}" for k, v in fabrications[:5]))
        except Exception as _fe:
            fabrications = []
            log.warning("Fabrication guard skipped for app %d (non-fatal): %s",
                        application_id, _fe)
        grounding_failed = False
        grounding_notes = None
        # Three states, not two: passed / failed / never ran. Collapsing the third
        # into "passed" is how an unchecked résumé came to be reported as verified.
        grounding_ran = False
        doctor_failed = False
        doctor_notes = None

        # 1. Grounding check — no hallucinated bullets
        try:
            from app.tailoring.grounding import GroundingChecker
            checker = GroundingChecker()
            log.info("Grounding check: app %d (variant: %s, attempt %d)...",
                     application_id, variant, attempt)
            g_result = checker.check(master, resume_md)
            # `unverified` is the checker saying it could not form an opinion
            # (no extractable bullets). That is the same state as "the check
            # never ran" — NOT a pass, and NOT a failure — so it flows into the
            # unverified path rather than shipping as verified.
            grounding_ran = not getattr(g_result, "unverified", False)
            grounding_tier = getattr(g_result, "tier", None)
            grounding_calls = int(getattr(g_result, "llm_calls", 0) or 0)
            grounding_cache_hits = int(getattr(g_result, "cache_hits", 0) or 0)
            grounding_spans_changed = int(getattr(g_result, "spans_changed", 0) or 0)
            grounding_spans_verified = int(getattr(g_result, "spans_verified", 0) or 0)
            if grounding_ran and not g_result.passed:
                grounding_failed = True
                grounding_notes = "Grounding check failed. Flagged bullets:\n" + "\n".join(
                    [f"- {fb['bullet']}" for fb in g_result.flagged_bullets]
                )
                log.warning("Grounding FAILED for app %d: %d bullets flagged",
                            application_id, len(g_result.flagged_bullets))
            else:
                log.info("Grounding PASSED for app %d (%s, %d/%d spans changed, "
                         "%d verified, %d cache hits, %d LLM calls)",
                         application_id, grounding_tier,
                         grounding_spans_changed,
                         int(getattr(g_result, "spans_total", 0) or 0),
                         grounding_spans_verified, grounding_cache_hits,
                         grounding_calls)
            # Feed the grounded ratio into the candidate's Trust Profile consistency
            # dimension (share of resume bullets supported by the master résumé).
            try:
                total = len(g_result.confidence_map) or 0
                if total:
                    ratio = max(0.0, (total - len(g_result.flagged_bullets)) / total)
                    from app.intelligence.trust_service import compute_and_store
                    compute_and_store(app_user_id, grounding_score=ratio)
            except Exception as _te:
                log.debug("trust consistency update skipped: %s", _te)
        except Exception as e:
            # The anti-hallucination gate did not run. This is an infrastructure
            # failure (grounding.py imports sentence_transformers at module level,
            # so a missing/broken ML stack or an unreachable model download lands
            # here), NOT a clean résumé — log it as such and let the honest
            # `grounding_ran=False` state flow through to the report and the UI.
            log.error("Grounding check DID NOT RUN for app %d — résumé is "
                      "unverified: %s", application_id, e)
            grounding_ran = False

        # 2. Resume Doctor — quality, ATS coverage, banned words, integrity
        try:
            from app.tailoring.doctor import ResumeDoctor
            doc = ResumeDoctor()
            d_result = doc.check(resume_md, master, job_description)
            doctor_score = d_result.score
            doctor_ats = d_result.ats_coverage_pct
            doctor_verdict = d_result.llm_verdict
            doctor_weak = d_result.weak_bullets or []
            doctor_banned = d_result.banned_found or []
            doctor_integrity = d_result.integrity_issues or []
            doctor_human = d_result.human_score
            doctor_fingerprints = d_result.fingerprint_flags or []
            # A SEPARATE outcome from doctor_failed. Reading machine-written is a
            # style defect, not a false claim, so it earns a rebuild but must
            # never park a truthful résumé at ERROR — the user would be left
            # with nothing to submit over a bolding count.
            human_failed = not getattr(d_result, "human_passed", True)
            log.info("Doctor report app %d: %s", application_id, d_result.summary())
            if not d_result.passed:
                doctor_failed = True
                doctor_notes = f"Doctor score={d_result.score}/100.\n" + "\n".join(d_result.issues)
                log.warning("Doctor FAILED for app %d (score=%d, attempt %d)",
                            application_id, d_result.score, attempt)
            elif human_failed:
                log.warning("Doctor: app %d reads machine-written (human=%d, attempt %d)",
                            application_id, d_result.human_score, attempt)
        except Exception as e:
            log.warning("Failed to run resume doctor: %s", e)

        # A fabricated fact is never shippable, and a rebuild is the only honest
        # response — so it drives the retry exactly like a grounding failure.
        fabrication_notes = None
        if fabrications:
            fabrication_notes = (
                "Facts appeared in the tailored résumé that are NOT in the master "
                "résumé. Remove or correct every one — use only what the master "
                "actually says:\n"
                + "\n".join(f"- invented {kind}: {value}" for kind, value in fabrications[:10])
            )

        if break_after_checks:
            break                      # L0: nothing was generated, so nothing to rebuild
        if not grounding_failed and not doctor_failed and not fabrications and not human_failed:
            break
        if attempt < MAX_TAILOR_ATTEMPTS:
            revision_notes = "\n".join(
                n for n in (grounding_notes, doctor_notes, fabrication_notes) if n)
            log.info("Rebuilding tailored resume for app %d (attempt %d failed review)",
                     application_id, attempt)

    # Build output paths — name files after the actual application owner.
    # Prefer the user's saved profile (multi-tenant); fall back to the static
    # QA store identity, then a generic name. Never use another user's name.
    first, last = "", ""
    if app_user_id:
        from app.db.models import UserProfile
        with get_session() as session:
            prof = session.exec(
                select(UserProfile).where(UserProfile.user_id == app_user_id)
            ).first()
            if prof:
                first = (prof.first_name or "").strip()
                last = (prof.last_name or "").strip()
    if not first and not last:
        identity = qa_resolver.data.get("identity", {})
        first = identity.get("first_name", "")
        last = identity.get("last_name", "")

    def _slug(s: str) -> str:
        return re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_")

    name_part = "_".join(p for p in (_slug(first), _slug(last)) if p) or "Candidate"
    resume_filename = f"{name_part}_Resume.docx"
    cover_filename = f"{name_part}_Cover_Letter.txt"

    out_dir = settings.data_dir / "tailored" / f"app_{application_id}"
    out_dir.mkdir(parents=True, exist_ok=True)
    resume_path = out_dir / resume_filename
    cover_path = out_dir / cover_filename
    _md_to_docx(resume_md, resume_path)
    # Keep the markdown the .docx was rendered from. Re-verifying a résumé by
    # parsing its own .docx back would check a lossy reconstruction rather than
    # the document we actually generated; the manual re-check below reads this.
    try:
        (out_dir / "resume.md").write_text(resume_md, encoding="utf-8")
    except Exception as _me:
        log.debug("tailored markdown not persisted: %s", _me)

    # Quality report for the Tailoring Studio UI (score dial, rebuilt badge,
    # keyword highlighting uses ats_keywords at read time).
    try:
        import json as _json
        (out_dir / "report.json").write_text(_json.dumps({
            "doctor_score": doctor_score,
            # A résumé that reads machine-written is NOT reported as passing.
            # Folding the anti-fingerprint result in here is the difference
            # between "we checked and it's fine" and "we checked, it isn't, and
            # we shipped it anyway with a green tick" — which is what a 24/100
            # human score used to get.
            "doctor_passed": (not doctor_failed) and not human_failed,
            # None (not True) when the check never ran — the UI must be able to
            # tell "verified clean" apart from "not verified".
            "grounding_passed": (not grounding_failed) if grounding_ran else None,
            "grounding_status": ("failed" if grounding_failed
                                 else "passed" if grounding_ran else "unverified"),
            "attempts": attempts_used,
            "variant": variant,
            # Rich, previously log-only quality feedback — surfaced in the UI so
            # the user sees WHY the score is what it is, not just the number.
            "ats_coverage_pct": round(doctor_ats * 100) if doctor_ats is not None else None,
            "verdict": doctor_verdict,
            "weak_bullets": doctor_weak[:5],
            "banned_words": doctor_banned[:8],
            "integrity_issues": doctor_integrity[:5],
            "human_score": doctor_human,
            "human_passed": not human_failed,
            "fingerprint_flags": doctor_fingerprints[:5],
            # Deterministic fact additions — employers, held titles, employment
            # dates, degrees, institutions, certifications, numbers. Empty is
            # the only acceptable value on a delivered résumé.
            "fabrications": [{"kind": k, "value": v} for k, v in fabrications[:10]],
            # What verification actually cost, so "we ground once per claim" is
            # a measurable statement rather than a design intention.
            "grounding_tier": grounding_tier,
            "grounding_llm_calls": grounding_calls,
            "grounding_cache_hits": grounding_cache_hits,
            "grounding_spans_changed": grounding_spans_changed,
            "grounding_spans_verified": grounding_spans_verified,
            "bold_spans_trimmed": bold_trimmed,
            # L0: the master already covered this posting, so nothing was rewritten.
            "tailoring_skipped": bool(skipped_reason),
            "tailoring_skipped_reason": skipped_reason,
            "master_coverage_pct": (round(master_coverage * 100)
                                    if master_coverage is not None else None),
            # Immutable facts the Lock layer pinned to the master résumé (shown
            # in the UI as "locked", so the user sees their degree/dates are
            # protected rather than seeing a scary "altered credential" flag).
            "locked_fields": locked_fields,
            "generated_at": datetime.utcnow().isoformat(),
        }), encoding="utf-8")
    except Exception as _re:
        log.debug("report.json write skipped: %s", _re)

    posted_str = job_posted_at.strftime("%B %d, %Y") if job_posted_at else "unknown"
    cover_header = (
        f"Company:   {job_company}\n"
        f"Role:      {job_title}\n"
        f"Posted:    {posted_str}\n"
        f"URL:       {job_url}\n"
        f"\n---COVER---\n\n"
    )
    cover_path.write_text(cover_header + cover, encoding="utf-8")

    # Mirror the tailored docs to durable storage (ephemeral local disk is wiped
    # on redeploy). Best-effort — never blocks the result write below.
    _persist_tailored_to_storage(
        app_user_id, application_id,
        [resume_path, cover_path, out_dir / "report.json", out_dir / "resume.md"])

    # --- Phase 3: write results in a short session ---
    with get_session() as session:
        app = session.get(Application, application_id)
        if not app:
            raise ValueError(f"Application {application_id} not found after LLM work")
        app.resume_variant = variant
        app.tailored_resume_path = str(resume_path)
        app.cover_letter_path = str(cover_path)
        app.updated_at = datetime.utcnow()

        _unverified_note = (
            "⚠ Grounding could not be verified — the anti-hallucination check did "
            "not run for this résumé (see server logs). Read every bullet against "
            "your master résumé before you submit."
        )
        _style_note = (
            "⚠ This draft still reads machine-written (uniform bullet lengths and "
            "openings). Every fact in it is verified against your master résumé — "
            "but vary a couple of bullets in your own words before you submit."
        )
        if fabrications:
            # The one failure with no benefit of the doubt: these facts are not in
            # the user's résumé, so they are not the user's facts. Nothing about a
            # deadline justifies handing someone a document that says they worked
            # somewhere they did not.
            app.status = ApplicationStatus.ERROR
            app.notes = (
                "Blocked: the tailored résumé asserted facts your master résumé "
                "does not contain — "
                + "; ".join(f"{kind} “{value}”" for kind, value in fabrications[:5])
                + ". Nothing was delivered. Rebuild, or edit your master résumé if "
                "these are real."
            )
            log.error("Application %d blocked at ERROR: %d fabricated fact(s)",
                      application_id, len(fabrications))
        elif grounding_failed:
            app.status = ApplicationStatus.ERROR
            app.notes = grounding_notes
            log.warning("Application %d blocked at ERROR: grounding failure", application_id)
        elif doctor_failed:
            app.status = ApplicationStatus.ERROR
            app.notes = doctor_notes
            log.warning("Application %d blocked at ERROR: doctor quality failure", application_id)
        elif not grounding_ran and settings.grounding_required:
            # Strict posture: refuse to hand over a résumé whose claims were never
            # checked. Off by default because a transient ML-stack failure would
            # otherwise take the whole tailoring feature down; on, this is the
            # fail-CLOSED behaviour.
            app.status = ApplicationStatus.ERROR
            app.notes = _unverified_note
            log.error("Application %d blocked at ERROR: grounding unverified and "
                      "GROUNDING_REQUIRED is set", application_id)
        else:
            app.status = ApplicationStatus.TAILORED
            notes = []
            if not grounding_ran:
                # Delivered, but never silently: the human review step is the
                # design's backstop and it can only work if the human is told.
                notes.append(_unverified_note)
            if human_failed:
                notes.append(_style_note)
            if skipped_reason:
                notes.append(skipped_reason)
            if notes:
                app.notes = "\n\n".join(notes)

        session.add(app)
        session.commit()

    return resume_path, cover_path


def reverify_application(application_id: int) -> dict:
    """Re-run verification on an already-delivered résumé, ignoring the cache.

    The cache exists because a fact-check is a pure function of (evidence,
    claim, verifier) — ask the same question twice and the honest answer is the
    stored one. But "the system is confident" is not the same as "the user is",
    and a candidate about to put their name on a document is entitled to make us
    look again. So this is the escape hatch: deterministic fact checks re-run
    from scratch, every needed verdict recomputed rather than read back.

    Exceptional by design. Nothing in the normal flow calls it, and the normal
    flow is not weaker for that — the automatic check already ran on this exact
    text before it was delivered.

    Returns the refreshed verification section of the report.
    """
    import json as _json
    from datetime import datetime

    with get_session() as session:
        app = session.get(Application, application_id)
        if not app:
            raise ValueError(f"Application {application_id} not found")
        app_user_id = app.user_id
        job = session.get(Job, app.job_id)
        job_description = job.description if job else ""

    out_dir = settings.data_dir / "tailored" / f"app_{application_id}"
    md_path = out_dir / "resume.md"
    if not md_path.exists():
        raise FileNotFoundError(
            "No tailored résumé markdown on file for this application — "
            "re-run tailoring before requesting a re-check."
        )
    resume_md = md_path.read_text(encoding="utf-8")

    from app.matching.pipeline import _load_resume
    master = _load_resume(user_id=app_user_id)

    from app.tailoring.evidence import fabrication_violations
    fabrications = fabrication_violations(master, resume_md)

    grounding_status = "unverified"
    flagged: list[str] = []
    calls = 0
    try:
        from app.tailoring.grounding import GroundingChecker
        result = GroundingChecker().check(master, resume_md, use_cache=False)
        calls = int(getattr(result, "llm_calls", 0) or 0)
        if not getattr(result, "unverified", False):
            grounding_status = "passed" if result.passed else "failed"
            flagged = [fb["bullet"] for fb in result.flagged_bullets]
    except Exception as e:
        log.error("Manual re-verification could not run for app %d: %s", application_id, e)

    doctor_ats = None
    try:
        from app.tailoring.doctor import ResumeDoctor
        doctor_ats = ResumeDoctor().check(resume_md, master, job_description or "").ats_coverage_pct
    except Exception as e:
        log.debug("re-verify doctor pass skipped for app %d: %s", application_id, e)

    section = {
        "reverified_at": datetime.utcnow().isoformat(),
        "grounding_status": grounding_status,
        "grounding_passed": (grounding_status == "passed") if grounding_status != "unverified" else None,
        "flagged_bullets": flagged[:10],
        "fabrications": [{"kind": k, "value": v} for k, v in fabrications[:10]],
        "grounding_llm_calls": calls,
        "ats_coverage_pct": round(doctor_ats * 100) if doctor_ats is not None else None,
    }

    report_path = out_dir / "report.json"
    try:
        existing = _json.loads(report_path.read_text(encoding="utf-8")) if report_path.exists() else {}
        existing.update(section)
        report_path.write_text(_json.dumps(existing), encoding="utf-8")
        _persist_tailored_to_storage(app_user_id, application_id, [report_path])
    except Exception as e:
        log.debug("re-verify report update skipped for app %d: %s", application_id, e)

    log.info("Manual re-verification of app %d: grounding=%s, %d fabrication(s), "
             "%d LLM call(s)", application_id, grounding_status, len(fabrications), calls)
    return section


def tailor_all_shortlisted(user_id: str | None = None) -> int:
    """Process every SHORTLISTED application."""
    with get_session() as session:
        q = select(Application).where(Application.status == ApplicationStatus.SHORTLISTED)
        if user_id:
            q = q.where(Application.user_id == user_id)
        apps = session.exec(q).all()
        ids = [a.id for a in apps]
    for aid in ids:
        try:
            tailor_for_application(aid)
            log.info("Tailored application %d", aid)
        except Exception as e:
            log.exception("Tailoring failed for app %d: %s", aid, e)
    return len(ids)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s: %(message)s")
    n = tailor_all_shortlisted()
    print(f"Tailored {n} applications.")
